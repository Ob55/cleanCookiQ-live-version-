"""Generate a non-technical Word document summarising the data calculations
and methodologies used throughout the cleancookIQ platform."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\USER\AD_DEV\cleanCookiQ-live-version-\cleancookIQ-data-methodologies.docx"

doc = Document()

# ---------- styling helpers ----------
def set_base_styles():
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

def H1(text):
    p = doc.add_heading(text, level=1)
    return p

def H2(text):
    p = doc.add_heading(text, level=2)
    return p

def H3(text):
    p = doc.add_heading(text, level=3)
    return p

def P(text=""):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def kv(label, value):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(value)
    return p

set_base_styles()

# ---------- title page ----------
title = doc.add_heading("cleancookIQ — How the Platform's Numbers Are Calculated", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A plain-English guide to every score, savings figure, "
                "carbon estimate and risk rating shown on the platform")
r.italic = True
r.font.size = Pt(12)

doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run(
    "cleancookIQ shows a lot of numbers — readiness scores, payback years, "
    "carbon tonnes avoided, deal-match percentages, risk ratings. "
    "This document explains, for each one: what it means, what goes in, "
    "what comes out, and where the underlying data comes from. "
    "It is written for programme managers, funders, partners and board "
    "members. No coding or finance background is assumed."
)

doc.add_paragraph()
note = doc.add_paragraph()
note_r = note.add_run(
    "The platform is organised internally into eight numbered \"workstreams\" "
    "(WS0 to WS8). You will see those labels referenced in screens; this "
    "document is grouped the same way so it is easy to map back."
)
note_r.italic = True

doc.add_page_break()

# ---------- section 0 : sourced numbers ----------
H1("1. Where every number comes from — the \"Sourced\" badge")

P("Every figure on cleancookIQ is meant to be defensible. Behind the scenes, "
  "the platform keeps a library of \"data points\" — individual values like "
  "\"the price of charcoal in Nakuru in March 2026\" or \"the CO₂ released per "
  "kilogram of firewood\". Each data point carries:")

bullet("the value and its unit (e.g. KSh 45 / kg)")
bullet("which county it applies to (or \"national\" if it's a country-wide figure)")
bullet("which fuel it relates to (firewood, charcoal, LPG, biogas, electric)")
bullet("the source — publisher, URL, publication date")
bullet("a confidence level: high, medium, modeled, or preliminary")
bullet("a validity window — when the value starts and stops being current")

P("When the platform needs a number, it follows a fixed priority: it first "
  "looks for a value that matches the user's exact county AND fuel; if there "
  "is none, it tries the county only, then a national value for that fuel, "
  "then a generic national fallback. The little info icon next to a number "
  "(the \"Sourced\" badge) lets anyone click through to see the citation, "
  "publisher and confidence level for that specific figure.")

P("This is the foundation everything else sits on. If a number on the "
  "platform looks wrong, the conversation should always start at the data "
  "point — not at the calculation that used it.")

# ---------- section 1 : readiness score ----------
H1("2. Institution Readiness Score (0 – 100)")

P("The Readiness Score is the single number that tells a transition manager "
  "how prepared a school, hospital or barracks is to switch to clean cooking "
  "today. It is calculated from eight things the institution tells us during "
  "onboarding:")

bullet("Their current cooking fuel (electric scores highest, firewood lowest)")
bullet("How much fuel they use per term")
bullet("Whether they have a dedicated kitchen at all")
bullet("The condition of that kitchen (clean-and-ready vs. needs major work)")
bullet("Their financing preference (open to a loan vs. grant-only)")
bullet("Number of students (bigger institutions score higher — economies of scale)")
bullet("Monthly fuel spend (higher spend means a bigger payoff to switching)")
bullet("Who actually makes the financial decision (a head teacher who can sign vs. waiting on the county)")

P("Each of those eight inputs gets a 0–100 sub-score from a lookup table or "
  "a banded threshold. The eight sub-scores are then averaged with these "
  "weights:")

kv("Kitchen exists", "20%")
kv("Kitchen condition", "20%")
kv("All other six factors", "10% each")

P("The final 0–100 number is grouped into four readiness categories:")

kv("80 – 100", "Ready Now")
kv("60 – 79",  "Ready with Minor Actions")
kv("40 – 59",  "Needs Enabling Support")
kv("Below 40", "Longer-Term Opportunity")

P("There is also an admin screen (\"Scoring Weight Configuration\") where "
  "the platform owner can adjust seven readiness dimensions. ")
warn = doc.add_paragraph()
wr = warn.add_run("Note: the admin weights screen is not yet wired into the "
                  "live calculation — the hard-coded weights above are what "
                  "currently drives the score on the platform.")
wr.italic = True

# ---------- section 2 : cooking cost calculator ----------
H1("3. \"Cooking Counting\" — savings calculator (institution view)")

P("This is the screen an institution sees when they want to know: \"if we "
  "switch to clean cooking, how much money and how much CO₂ would we save "
  "each year?\" It assumes 3 school terms per year and uses these inputs:")

bullet("How much fuel they use per term, in their unit (kg, bag, m³…)")
bullet("The cost per unit of that fuel, looked up from the Sourced data layer for their county")
bullet("The CO₂ released per unit of that fuel (also Sourced)")
bullet("A \"clean cost multiplier\" — typically 0.4, meaning clean cooking costs about 40% of the firewood cost")
bullet("A \"CO₂ reduction fraction\" — typically 0.85, meaning the switch cuts emissions by about 85%")
bullet("A \"cooking time savings fraction\" — typically 0.5, meaning kitchen staff save about half their cooking time")

P("The screen then shows:")

kv("Current annual cost", "consumption per term × cost per unit × 3 terms")
kv("Annual cost after switching", "current cost × clean multiplier")
kv("Annual savings", "the difference between the two")
kv("CO₂ avoided per year", "current emissions × CO₂ reduction fraction")
kv("Cooking time saved per year", "minutes per day × 0.5 × 270 cooking days, shown in hours")
kv("5-year comparison chart", "current cost line vs. clean-cooking cost line, year by year")

P("Every one of those numbers carries the Sourced badge so the school can "
  "click and see exactly where the cost-per-kilo, CO₂ factor or savings "
  "fraction came from.")

# ---------- section 3 : TCO / financing model ----------
H1("4. Total Cost of Ownership (TCO) and financing model")

P("This is the heart of the financing screens. It answers the question: "
  "\"if a hospital buys a biogas system today, will it actually pay back, "
  "and over how many years?\" Inputs include the equipment cost, install "
  "cost, expected lifetime, maintenance, salvage value at end-of-life, the "
  "fuel cost the institution would have paid otherwise, and the structure "
  "of any grant, loan or down-payment.")

H3("What the model produces")

bullet("A year-by-year cash flow table: outflows for capex, opex, maintenance and loan repayments; inflows for the grant, salvage and the fuel cost they no longer have to pay.")
bullet("Net Present Value (NPV) — the value of the whole project today, expressed in shillings, after discounting future years at a chosen rate (default 12%). A positive NPV means the project is worth doing on financial terms alone.")
bullet("Internal Rate of Return (IRR) — the implicit \"interest rate\" the project earns. Funders compare this to their own cost of capital.")
bullet("Simple payback — the year at which cumulative savings cross zero.")
bullet("Discounted payback — the same, but accounting for the time-value of money.")

H3("Built-in assumptions")

bullet("Operating costs rise about 5% per year (fuel inflation)")
bullet("Maintenance rises about 3% per year")
bullet("The status-quo (firewood/charcoal) cost rises about 7% per year — this is what makes clean cooking look better over time")
bullet("Default install cost is 10% of equipment cost; default loan rate 8%; default tenor 36 months")

H3("Sensitivity analysis")

P("Because every assumption can be wrong, the model also runs three \"what-"
  "if\" scenarios automatically and shows how NPV and payback change when:")

bullet("Operating cost is 20% higher or lower than expected")
bullet("Equipment lifetime is 20% shorter or longer than expected")
bullet("The discount rate is 5 percentage points higher or lower")

P("This lets a funder see at a glance: \"is this project still bankable if "
  "fuel inflation disappoints?\"")

H3("Where the model is used")

bullet("Admin > Financing Designer — for building reference deals")
bullet("Funder > Deal Flow — to surface NPV/IRR per opportunity")
bullet("Institution dashboards — for the headline payback figure")

# ---------- section 4 : marketplace / CSCC ----------
H1("5. Supplier Compliance Tier (CSCC)")

P("The marketplace tells buyers which suppliers are safe to procure from. "
  "Each supplier fills in a Clean Stove Compliance Checklist (CSCC) covering "
  "six sections: A (general business compliance), B (cooking equipment), "
  "C (installation), D (biogas), E (LPG), and F (self-declarations).")

P("The platform reads those tick-boxes and assigns a tier:")

kv("Tier 3 — Uncertified", "Supplier has self-declared they are not certified. Not recommended for procurement.")
kv("Tier 2 — In Progress", "Supplier has a KEBS application pending. Provisional listing only.")
kv("Tier 1 — Compliant",   "Section A is fully complete AND at least one of B/C/D/E is fully complete (i.e. the supplier has fully declared at least one product line).")
kv("Unrated",              "No CSCC submission yet. No compliance signal available.")

P("The tier badge appears on every supplier card and product listing.")

H2("Marketplace filters and ratings")
P("Products can be filtered by category, county served, in-stock-only, and "
  "price range. Average ratings are computed from published buyer reviews "
  "(simple mean, rounded to two decimals).")

# ---------- section 5 : delivery ----------
H1("6. Delivery and installation tracking")

P("Once an institution has chosen a supplier, the platform tracks the "
  "delivery through nine sequential stages:")

stages = ["Manufacturing", "Dispatched", "In transit", "On site",
          "Installing", "Commissioned", "Handover", "Acceptance window",
          "Monitoring"]
for s in stages:
    bullet(s)

P("(Plus a \"Cancelled\" state that can be reached from anywhere.)")

P("A simple progress percentage is shown by counting how many stages are "
  "complete: e.g. reaching \"Commissioned\" means 6 of 9, so 67% progress.")

P("Each commissioning has a checklist of physical-acceptance items — power "
  "tested, gas connected, staff trained, etc. The percent-complete on that "
  "checklist drives whether the institution can sign off the delivery. "
  "Sign-off is only allowed once the checklist is 100% complete AND the "
  "stage has reached at least \"Handover\".")

# ---------- section 6 : risk ----------
H1("7. Risk Register and scoring")

P("Every project has a risk register. Each risk is rated on two 1–5 scales:")

kv("Severity", "how bad it would be if it happened (1 = trivial, 5 = catastrophic)")
kv("Likelihood", "how probable it is (1 = very unlikely, 5 = expected)")

P("The platform multiplies them together for a 1–25 score, then assigns a band:")

kv("16 – 25", "Critical (red)")
kv("9 – 15",  "High (amber)")
kv("4 – 8",   "Medium (yellow)")
kv("1 – 3",   "Low (green)")

P("Risk types tracked include equipment defects, installation failures, "
  "fuel-price spikes, demand drops, behavioural relapse, currency / import "
  "shocks, counterparty closure, carbon non-issuance, cybersecurity, and "
  "regulatory.")

H2("Behavioural relapse — the automatic tripwire")

P("\"Relapse\" means an institution has been given clean cookers but has "
  "drifted back to firewood or charcoal. The platform watches for this "
  "automatically.")

P("Every monitoring reading records two numbers: how much clean fuel was "
  "used, and how much firewood/charcoal was still being used. From these we "
  "compute the \"clean fuel share\":")

kv("Clean fuel share", "clean fuel used ÷ (clean fuel + baseline fuel)")

P("If the most recent two readings both have a clean-fuel share below 50%, "
  "the platform automatically:")

bullet("Opens a \"Behavioural relapse\" risk on the project at severity 4 × likelihood 4 = score 16 (Critical)")
bullet("Opens a high-priority support ticket titled \"Refresher training required (relapse detected)\"")
bullet("Both actions are idempotent — they will not duplicate if one is already open")

P("This is the platform's main safety net for ensuring carbon and impact "
  "claims stay honest.")

# ---------- section 7 : carbon ----------
H1("8. Carbon credits and verification")

P("For projects pursuing carbon credits, the platform tracks three layers "
  "of tonnage:")

kv("Estimated annual credits", "computed at design time from a single rough formula")
kv("Total estimated tCO₂e",    "sum of period-by-period forecasts in the credit-estimates ledger")
kv("Total verified tCO₂e",     "sum of independently-audited verifications by an accredited Validation/Verification Body (VVB) such as SCS or TÜV SÜD")

P("The rough annual estimate is calculated as:")

f = doc.add_paragraph()
fr = f.add_run("(baseline fuel × baseline CO₂ factor − project fuel × project "
               "CO₂ factor) × periods per year ÷ 1000")
fr.italic = True
fr.font.size = Pt(10)

P("…which gives tonnes of CO₂ avoided per year, never below zero. The "
  "factors come from the same Sourced data layer described in section 1 "
  "(the impact-report screen cites IPCC AR6 emission factors as the "
  "intended source).")

P("Each carbon project also tracks: methodology (e.g. Gold Standard "
  "TPDDTEC v3.1), registry (Gold Standard, Verra, Plan Vivo), registry "
  "project ID, vintage, and status (design → validation → registered → "
  "issued → retired, or rejected).")

# ---------- section 8 : funder portal ----------
H1("9. Funder Portal — deal matching, portfolio and impact attribution")

H2("Deal-match score (0% – 100%)")

P("When a funder sets their preferences (counties, fuels, ticket-size band, "
  "max acceptable risk score, minimum IRR), the platform scores every open "
  "project against those preferences. The score has five components, with "
  "fixed weights:")

kv("County preference match",  "30%  (binary — the project's county is, or isn't, on their list)")
kv("Fuel preference match",    "20%  (binary)")
kv("Ticket-size fit",          "25%  (1.0 if the funding gap sits inside their min/max band; tapers off linearly outside)")
kv("Risk fit",                 "15%  (1.0 if the worst open risk is below their cap; degrades the further over the cap it goes)")
kv("IRR fit",                  "10%  (placeholder — currently 1.0 unless they set an IRR floor and we have no project IRR yet, in which case 0.5)")

P("If a project misses a hard preference (wrong county, wrong fuel, or "
  "risk far outside their cap), the total is capped at 20% so it sinks to "
  "the bottom of their list.")

H2("Funding gap")
P("On every project: total budget − money already committed or disbursed by "
  "any funder. This is what other funders are being invited to fill.")

H2("Portfolio summary")
P("For each funder we show:")

bullet("Number of projects supported")
bullet("Total capital committed and disbursed")
bullet("Lifetime tCO₂e, KSh savings, meals served, and jobs created — attributable to them")

H2("How attribution works")
P("When several funders co-fund one project, the platform splits the "
  "outcomes between them in proportion to the capital each one put in. "
  "If Funder A contributed 60% of the capital and Funder B 40%, then 60% "
  "of the tCO₂e avoided is credited to A and 40% to B. This is the "
  "methodology cited in the auto-generated quarterly impact report.")

# ---------- section 9 : county intelligence ----------
H1("10. County Intelligence")

P("The county pages roll up everything happening in one of Kenya's 47 "
  "counties:")

bullet("Number of institutions on the platform in that county")
bullet("How many have been assessed or further along the pipeline")
bullet("How many have actually transitioned (installed or already monitoring)")
bullet("The dominant baseline fuel in that county (the most common current cooking fuel)")
bullet("Total meals per day and total students across institutions there")
bullet("Number of suppliers that serve the county")
bullet("Latest fuel prices observed in that county for each fuel type")
bullet("Active local policies that affect clean cooking")

P("These figures recompute live from the underlying tables, so they are "
  "always current.")

# ---------- section 10 : pipeline & portfolio aggregation ----------
H1("11. Pipeline funnel and portfolio aggregation (admin)")

P("The admin Pipeline Dashboard shows three stages:")

bullet("Identified — institutions on the books")
bullet("Assessed / Scored — institutions whose readiness score has been computed")
bullet("Installed — equipment is in")

P("Each stage shows a count and what percentage of all institutions it "
  "represents, so it is easy to see where the funnel is narrowing.")

P("The Portfolio Aggregation screen lets an admin group institutions into "
  "named portfolios (e.g. \"Eastern Province secondary schools\") and shows "
  "the totals for each group:")

bullet("Number of institutions")
bullet("Total annual KSh savings (sum of each institution's stored figure)")
bullet("Total tonnes of CO₂ reduction per year (same)")

# ---------- section 11 : onboarding wizard ----------
H1("12. Onboarding wizard logic")

P("The multi-step setup wizards (institution onboarding, funder onboarding, "
  "supplier onboarding) all use the same engine. For each step we know:")

bullet("Whether the user can move forward (the step is filled in correctly)")
bullet("Whether the user can move back")
bullet("Validation issues to display (e.g. \"email must be valid\")")
bullet("The progress bar — what fraction of the wizard is complete")

P("Validators are simple, composable checks: required fields, minimum "
  "length, valid email format, etc. Nothing more elaborate.")

# ---------- section 12 : known gaps ----------
H1("13. Known gaps in the methodology — items to close")

P("In the spirit of full disclosure, here are places where the platform's "
  "calculations are not yet wired end-to-end. None of these break what is "
  "shown today, but they all matter for full defensibility:")

H3("a) Admin readiness weights are not yet live")
P("The Scoring Weight Configuration screen lets an admin tune seven "
  "dimensions, but the actual readiness score still uses hard-coded weights. "
  "Until the two are connected, changes in the admin screen have no effect "
  "on the score shown to users.")

H3("b) Two parallel cost models")
P("There are two separate ways the platform stores technology costs: a "
  "per-student capex/opex/CO₂ table (Cost Configuration), and a per-"
  "technology profile with capex range, lifetime and efficiency (Financing "
  "Designer). Nothing reconciles them. A future cleanup should pick one "
  "as the source of truth.")

H3("c) Stored savings figures are entered manually")
P("On every institution we store \"annual savings (KSh)\", \"CO₂ reduction "
  "(t/yr)\" and \"recommended solution\". These drive the portfolio "
  "dashboards. Right now they are filled in by hand — they are not "
  "auto-derived from the TCO model and the carbon estimator. A future "
  "service should recompute them whenever the inputs change.")

H3("d) IRR fit in deal matching is a placeholder")
P("The 10% IRR-fit weight in the deal-match score currently always "
  "returns 1.0 (or 0.5 if the funder set a floor but we have no project "
  "IRR yet). Once project-level IRR from the TCO model flows into the "
  "funder deal flow view, this becomes a real signal.")

H3("e) Carbon project totals don't roll up monitoring readings yet")
P("The baseline and project emissions stored on each carbon project are "
  "independent of the per-period monitoring readings. There is no service "
  "that aggregates monitoring readings into the carbon project totals "
  "automatically. Verifications still come from the audit ledger, which is "
  "the right answer for issued credits — but the forecast totals would be "
  "more credible if they pulled live from monitoring.")

H3("f) Marketing analysis page numbers are hard-coded")
P("The headline numbers on the public Market Insights page (\"340% growth\", "
  "\"2,847 institutions\") are static text. They do not come from the live "
  "county-metrics view. Either they should be wired up, or marked clearly "
  "as published-report figures with a date.")

H3("g) Sensitivity analysis covers only three drivers")
P("The TCO sensitivity table flexes operating cost, equipment lifetime and "
  "discount rate. It does not flex capital cost, baseline fuel cost or "
  "fuel-price escalation — the three drivers that, in practice, push "
  "payback most. Adding them is a small change with a high-value payoff "
  "for funder confidence.")

# ---------- closing ----------
doc.add_page_break()
H1("In summary")

P("Every number on cleancookIQ ladders up from a small, auditable set of "
  "rules:")

bullet("Sourced data points provide the raw inputs (prices, factors, validity).")
bullet("Pure formulas — readiness, savings, NPV, IRR, payback, deal-match, attribution, relapse — turn those inputs into the figures users see.")
bullet("Database views aggregate across institutions, counties, projects and funders for the dashboards.")
bullet("Triggers and watchdogs (relapse detection, sign-off rules) keep impact claims honest after delivery.")

P("Anyone questioning a number on the platform should be able to walk back "
  "from the screen → to the formula → to the data point → to the original "
  "publisher. That is the design intent. The seven gaps listed above are "
  "the work needed to make that promise complete.")

doc.save(OUT)
print(f"Wrote: {OUT}")
